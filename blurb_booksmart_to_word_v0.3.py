'''
	Created by Ben Hooper
	https://mythofechelon.co.uk/
	
	v0.3
	
	Changes in v0.3 (2021/05/26): Fixed handling of invalid file path (was using "and" when should have been using "or", improved handling of invalid file path to it keeps re-propmpting until valid input is given instead of exiting the script (therefore removed importing of "sys"), implemented proper HTML unescaping, clarified that second prompt is for footer and must include leading and trailing spaces, 
	Changes in v0.2 (2021/05/22): Added prompting for file path and book title, improved output (readability and information on formatting applied).
	Changes in v0.1 (2021/05/09): Created.

'''

from docx import Document # Install using "pip install python-docx"
import defusedxml.ElementTree as ET # Install using "pip install defusedxml"
import re, os, html
import argparse

CLIParser = argparse.ArgumentParser(description='A program to convert blurp\'s *.BOOK file to DocX format.')

CLIParser.add_argument("-p", "--Path", help="Enter path to .BOOK file", type=str)
CLIParser.add_argument("-t", "--Title", help="Enter book footer / title (including leading or trailing spaces)", type=str)
CLIParser.add_argument("-lp", "--LogPath", help="Enter path for logfile. Include file path, filename, and file extention", type=str)
CLIParser.add_argument("-dp", "--DocPath", help="Enter path for docx file. Include file path, filename, and file extention", type=str)

CLIArgs = CLIParser.parse_args()

if not CLIArgs.Path == "":
    bookfile_original_path_full = CLIArgs.Path
if not CLIArgs.Title == "":
    bookfile_title = CLIArgs.Title
if not CLIArgs.LogPath == "":
     bookfile_log_path_full = CLIArgs.LogPath
if not CLIArgs.DocPath == "":
     bookfile_docx_path_full = CLIArgs.DocPath

bookfile_original_path_valid = False
while bookfile_original_path_valid == False:
    if bookfile_original_path_full == None:
        bookfile_original_path_full = input("Enter path to .BOOK file: ")
    bookfile_original_path_full = bookfile_original_path_full.strip("'").strip('"')
    bookfile_original_ext = bookfile_original_path_full.rsplit(".", 1)[1]
    if not os.path.exists(bookfile_original_path_full) or not bookfile_original_ext == "book":
        print("File path invalid.")
    else:
        bookfile_original_path_valid = True
bookfile_xml_path_full = bookfile_original_path_full + ".xml"

if bookfile_docx_path_full == None:
    bookfile_docx_path_full = bookfile_original_path_full + ".docx"
if bookfile_docx_path_full == None:
    bookfile_log_path_full = bookfile_original_path_full + ".log"

if bookfile_original_path_full == "":
    bookfile_title = input("Enter book footer / title (including leading or trailing spaces): ")

print("")

with open(bookfile_original_path_full, "r", encoding="utf-8", errors="replace") as bookfile_original_object:
    bookfile_original_content = bookfile_original_object.read()

# Blurb BookSmart's .book files use (1) an XML structure but a sort of HTML encoding (XHTML?) like "&amp;apos;" and "&amp;quot;" and (2) multiple namespacing statements
# Basic structure for sentences:
'''
  &lt;void method="add"&gt;
   &lt;object class="java.util.LinkedList"&gt;                ← Contains all textformatting and text
    &lt;void method="add"&gt;
     &lt;object class="java.util.HashMap"&gt;                 ← Contains all textformatting (number of children varies)
      &lt;void method="put"&gt;
       &lt;string&gt;resolver&lt;/string&gt;
       &lt;string&gt;bod_l_9-12_s4.chars&lt;/string&gt;
      &lt;/void&gt;
      &lt;void method="put"&gt;
       &lt;string&gt;size&lt;/string&gt;
       &lt;int&gt;12&lt;/int&gt;
      &lt;/void&gt;
      &lt;void method="put"&gt;
       &lt;string&gt;underline&lt;/string&gt;
       &lt;boolean&gt;true&lt;/boolean&gt;
      &lt;/void&gt;
      &lt;void method="put"&gt;
       &lt;string&gt;bold&lt;/string&gt;
       &lt;boolean&gt;true&lt;/boolean&gt;
      &lt;/void&gt;
      &lt;void method="put"&gt;
       &lt;string&gt;family&lt;/string&gt;
       &lt;string&gt;Times New Roman&lt;/string&gt;
      &lt;/void&gt;
      &lt;void method="put"&gt;
       &lt;string&gt;italic&lt;/string&gt;
       &lt;boolean&gt;true&lt;/boolean&gt;
      &lt;/void&gt;
     &lt;/object&gt;
    &lt;/void&gt;
    &lt;void method="add"&gt;                                 ← Contains text that the above text formatting applies to
     &lt;string&gt;Think carefully about that&lt;/string&gt;
    &lt;/void&gt;
   &lt;/object&gt;
  &lt;/void&gt;
'''
# The next 2 lines correct this
xml_namespacingstatement = '<?xml version="1.0" encoding="UTF-8"?>'
bookfile_xml_content = html.unescape(bookfile_original_content.replace("&amp;", "&"))
bookfile_xml_content = bookfile_xml_content.replace(xml_namespacingstatement, "")
bookfile_xml_content = xml_namespacingstatement + bookfile_xml_content

# Create .xml file version of .book file
with open(bookfile_xml_path_full, "w", encoding="utf-8") as bookfile_xml_object:
    bookfile_xml_object.write(bookfile_xml_content)

# Parse XML so the nodes can be navigated
bookfile_xml_ET_root = ET.fromstring(bookfile_xml_content)
bookfile_xml_ET_nodes_linkedlist = bookfile_xml_ET_root.findall('.//object[@class="java.util.LinkedList"]')

# Prepare the Word document
document = Document()

# Create accompanying log file
bookfile_log_object = open(bookfile_log_path_full, "w", encoding="utf-8")

lastParagraphSingleLetter = False
for counter, node_linkedlist in enumerate(bookfile_xml_ET_nodes_linkedlist):
    newparagraphneeded = False
    
    node_linkedlist_children = list(node_linkedlist)
    
    textformatting_italic = False
    textformatting_bold = False
    textformatting_underline = False
    node_formatting = node_linkedlist_children[0][0] # Equivalent to <object class="java.util.HashMap">
    for node_void_put in node_formatting:
        formatting_label = node_void_put[0].text
        formatting_value = node_void_put[1].text
        if formatting_label == "italic":
            textformatting_italic = True
        if formatting_label == "bold":
            textformatting_bold = True
        if formatting_label == "underline":
            textformatting_underline = True
    
    node_string = node_linkedlist_children[1][0] # Equivalent to <void method="add"><string>
    string = node_string.text
    # The next line / if statement looks for paragraph indentations - strings that start with one tab or more, 2 spaces or more, or 2 tabs or spaces or more. The very first few lines look like this so we blanket allow those.
    if re.search("^(\t+|\s{2,}|[\t\s]{2,})[^\s]+", string) or counter < 5:
        newparagraphneeded = True
    
    # Check if string only contains a single character. If so, no new paragraph needed. 
    # A single character usually means a single big letter that prefix a word. So it's better to just make it a single word in the converting process.
    if len(string.strip()) == 1:
        lastParagraphSingleLetter = True
        newparagraphneeded = True

    if lastParagraphSingleLetter:
        newparagraphneeded = False
        lastParagraphSingleLetter = False

    # The next line / if statement looks for and excludes header or footer content (book title and page numbers) and empty lines
    if (string != bookfile_title) and (not re.search("^\d{1,4}$", string)) and (not re.search("^\n\s*$", string)):
        print("")
        print("--------------------------------------------------")
        print("")
        
        print("Current raw string:\n'{}'\n".format(string))
        bookfile_log_object.write("Current raw string:\n'{}'\n".format(string))
        
        if newparagraphneeded == True:
            print("New paragraph detected as being needed. Adding...")
            bookfile_log_object.write("New paragraph detected as being needed. Adding...\n")
            paragraph = document.add_paragraph()
        
        if string.endswith("\n"):
            print("Trailing newline detected. Removing...")
            bookfile_log_object.write("Trailing newline detected. Removing...\n")
            string = string[:-1]
        
        '''
        if string.startswith("Chapter"):
            document.add_heading(string, 1)
        '''

        # The first proper string sometimes does not have a paragraph to write too. So we create one here.
        while True:
            try:
                run = paragraph.add_run(string)
            except NameError:
                print("New paragraph detected as being needed. Adding...")
                bookfile_log_object.write("New paragraph detected as being needed. Adding...\n")
                paragraph = document.add_paragraph()
                continue
            break

        font = run.font
        if textformatting_italic == True:
            print("Italic formatting detected. Adding...")
            bookfile_log_object.write("Italic formatting detected. Adding...\n")
            font.italic = True
        if textformatting_bold == True:
            print("Bold formatting detected. Adding...")
            bookfile_log_object.write("Bold formatting detected. Adding...\n")
            font.bold = True
        if textformatting_underline == True:
            print("Underline formatting detected. Adding...")
            bookfile_log_object.write("Underline formatting detected. Adding...\n")
            font.underline = True
            
        print("\nCurrent string complete. Moving on to the next...")
        bookfile_log_object.write("\nCurrent string complete. Moving on to the next...\n")
    
document.save(bookfile_docx_path_full)
bookfile_log_object.close()