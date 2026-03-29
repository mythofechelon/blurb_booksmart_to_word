"""
Description:
	Take a .BOOK file from Blurb BookSmart and convert it to a .DOCX Microsoft Word file.

Pre-GitHub version history:
	• v0.3:
		- Author: Ben Hooper (https://mythofechelon.co.uk/)
		- Date: 2021/05/26
		- Changes:
			= Fixed handling of invalid file path.
			= Improved handling of invalid file path so it keeps re-propmpting until valid input is given, instead of exiting the script.
			= Implemented proper HTML unescaping.
			= Clarified that second prompt is for footer and must include leading and trailing spaces.
	• v0.2:
		- Author: Ben Hooper (https://mythofechelon.co.uk/)
		- Date: 2021/05/22
		- Changes:
			= Added prompting for file path and book title.
			= Improved output (readability and information on formatting applied).
	• v0.1:
		- Author: Ben Hooper (https://mythofechelon.co.uk/)
		- Date: 2021/05/09
		- Changes:
			= Base functionality.

Improvement opportunities:
	• None known.
"""



from pathlib import Path

import re

import os

import html

from sys import platform

if platform == "linux" or platform == "darwin":
	shell = "Terminal"
elif platform == "win32":
	shell = "Command Prompt or PowerShell"
else:
	shell = "a shell / terminal"

package_install_instructions = "Required package not installed. To fix this:\n1. Search for and open {shell}.\n2. Type the following and press enter: pip install {package_name}"

try:
	from docx import Document
except ModuleNotFoundError:
	package_name = "python-docx"
	
	print(package_install_instructions.format(shell=shell, package_name=package_name))
	
	exit()

try:
	import defusedxml.ElementTree as ET
except ModuleNotFoundError:
	package_name = "defusedxml"
	
	print(package_install_instructions.format(shell=shell, package_name=package_name))
	
	exit()


def log(message, log_file):
	print(message)
	log_file.write(message)

def main():
	while True:
		book_file_path = Path(input("\nEnter the full path to the .BOOK file.\n").strip("'\" "))
		
		if not book_file_path.exists():
			print("\n[ERROR] File at given path does not exist.")
			continue
			
		elif not book_file_path.is_file():
			print("\n[ERROR] Given path is not for a file.")
			continue
			
		elif not book_file_path.suffix.lower() == ".book":
			print("\n[ERROR] Given path is not for a .BOOK file.")
			continue
			
		break
		
	xml_file_path = str(book_file_path) + ".xml"
	docx_file_path = str(book_file_path) + ".docx"
	log_file_path = str(book_file_path) + ".log"

	book_title = input("\nEnter the book footer / title, including any leading or trailing spaces. If this is not known, simply press enter and this script will try to locate and extract it from the data.\n")

	with open(book_file_path, "r", encoding="utf-8", errors="replace") as book_file:
		book_file_content = book_file.read()

	# Blurb BookSmart's .book files use (1) an XML structure but a sort of HTML encoding (XHTML?) like "&amp;apos;" and "&amp;quot;" and (2) multiple namespacing statements
	# Basic structure for sentences:
	'''
	  &lt;void method="add"&gt;
	   &lt;object class="java.util.LinkedList"&gt;                ? Contains all textformatting and text
		&lt;void method="add"&gt;
		 &lt;object class="java.util.HashMap"&gt;                 ? Contains all textformatting (number of children varies)
		  &lt;void method="put"&gt;
		   &lt;string&gt;resolver&lt;/string&gt;
		   &lt;string&gt;bod_l_9-12_s4.chars&lt;/string&gt;
		  &lt;/void&gt;
		  &lt;void method="put"&gt;
		   &lt;string&gt;size&lt;/string&gt;
		   &lt;int&gt;12&lt;/int&gt;
		  &lt;/void&gt;
		  &lt;void method="put"&gt;
		   &lt;string&gt;underline&lt;/string&gt;
		   &lt;boolean&gt;true&lt;/boolean&gt;
		  &lt;/void&gt;
		  &lt;void method="put"&gt;
		   &lt;string&gt;bold&lt;/string&gt;
		   &lt;boolean&gt;true&lt;/boolean&gt;
		  &lt;/void&gt;
		  &lt;void method="put"&gt;
		   &lt;string&gt;family&lt;/string&gt;
		   &lt;string&gt;Times New Roman&lt;/string&gt;
		  &lt;/void&gt;
		  &lt;void method="put"&gt;
		   &lt;string&gt;italic&lt;/string&gt;
		   &lt;boolean&gt;true&lt;/boolean&gt;
		  &lt;/void&gt;
		 &lt;/object&gt;
		&lt;/void&gt;
		&lt;void method="add"&gt;                                 ? Contains text that the above text formatting applies to
		 &lt;string&gt;Think carefully about that&lt;/string&gt;
		&lt;/void&gt;
	   &lt;/object&gt;
	  &lt;/void&gt;
	'''
	# The next 2 lines correct this
	xml_namespacingstatement = '<?xml version="1.0" encoding="UTF-8"?>'
	xml_file_content = html.unescape(book_file_content.replace("&amp;", "&"))
	xml_file_content = xml_file_content.replace(xml_namespacingstatement, "")
	xml_file_content = xml_namespacingstatement + xml_file_content

	# Create .xml file version of .book file
	with open(xml_file_path, "w", encoding="utf-8") as xml_file_object:
		xml_file_object.write(xml_file_content)

	# Parse XML so the nodes can be navigated
	xml_file_ET_root = ET.fromstring(xml_file_content)
	
	if not book_title:
		# Extract book title
		for text_content in xml_file_ET_root.iter("TextContent"):
			if text_content.get("dc") != "$BookTitle":
				continue

			outer_list = text_content.find("./dm/java/object[@class='java.util.LinkedList']")
			if outer_list is None:
				continue

			outer_adds = outer_list.findall("./void[@method='add']")
			if len(outer_adds) < 2:
				continue

			# First outer add = metadata hashmap
			resolver_elem = outer_adds[0].find(
				"./object[@class='java.util.HashMap']"
				"/void[string='resolver']/string[2]"
			)

			if resolver_elem is None or resolver_elem.text != "Header/Footer.chars":
				continue

			# Second outer add = actual displayed text list
			title_elem = outer_adds[1].find(
				"./object[@class='java.util.LinkedList']"
				"/void[@method='add'][2]/string"
			)

			if title_elem is not None:
				book_title = title_elem.text
				break
	
	xml_file_ET_nodes_linkedlist = xml_file_ET_root.findall('.//object[@class="java.util.LinkedList"]')

	# Prepare the Word document
	document = Document()

	# Create accompanying log file
	log_file = open(log_file_path, "w", encoding="utf-8")

	for counter, node_linkedlist in enumerate(xml_file_ET_nodes_linkedlist):
		new_paragraph_needed = False
		
		node_linkedlist_children = list(node_linkedlist)
		
		textformatting_italic = False
		
		textformatting_bold = False
		
		textformatting_underline = False
		
		node_formatting = node_linkedlist_children[0][0] # Equivalent to <object class="java.util.HashMap">
		
		for node_void_put in node_formatting:
			formatting_label = node_void_put[0].text
			
			formatting_value = node_void_put[1].text
			
			if formatting_label == "italic":
				textformatting_italic = True
				
			if formatting_label == "bold":
				textformatting_bold = True
				
			if formatting_label == "underline":
				textformatting_underline = True
		
		node_string = node_linkedlist_children[1][0] # Equivalent to <void method="add"><string>
		string = node_string.text
		
		# The next line / if statement looks for paragraph indentations - strings that start with one tab or more, 2 spaces or more, or 2 tabs or spaces or more. The very first few lines look like this so we blanket allow those.
		if re.search(r"^(\t+|\s{2,}|[\t\s]{2,})[^\s]+", string) or counter < 5:
			new_paragraph_needed = True
		
		# The next line / if statement looks for and excludes header or footer content (book title and page numbers) and empty lines
		if (string != book_title) and (not re.search(r"^\d{1,4}$", string)) and (not re.search(r"^\n\s*$", string)):
			log(f"\nCurrent raw string:\n'{string}'", log_file)
			
			if new_paragraph_needed:
				log("\nNew paragraph detected as being needed. Adding...", log_file)
				
				paragraph = document.add_paragraph()
			
			if string.endswith("\n"):
				log("\nTrailing newline detected. Removing...", log_file)
				
				string = string[:-1]
			
			"""
			if string.startswith("Chapter"):
				document.add_heading(string, 1)
			"""
			
			run = paragraph.add_run(string)
			
			font = run.font
			
			if textformatting_italic == True:
				log("\nItalic formatting detected. Adding...", log_file)
				
				font.italic = True
				
			if textformatting_bold == True:
				log("\nBold formatting detected. Adding...", log_file)
				
				font.bold = True
				
			if textformatting_underline == True:
				log("\nUnderline formatting detected. Adding...", log_file)
				
				font.underline = True
				
			print("\n--------------------------------------------------")
		
	document.save(docx_file_path)
	
	log_file.close()
	
	print("\n[Sucess] File converted. Opening...")
	
	os.startfile(docx_file_path)

if __name__ == "__main__":
	main()